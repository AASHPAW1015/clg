import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def create_document(data_list, output_filename="HTML_CSS_Differentiation_Version2.docx"):
    doc = docx.Document()

    # Document properties
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'  # Changed font for version 2
    font.size = Pt(11)

    # Title
    title = doc.add_heading('HTML & CSS Differentiation - V2', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Space

    current_section = ""

    for item in data_list:
        section = item.get('section', '')
        if section != current_section:
            # Add section heading
            sec_heading = doc.add_heading(section, level=1)
            sec_heading.style.font.color.rgb = RGBColor(34, 139, 34) # Changed color to Forest Green
            current_section = section
            doc.add_paragraph()

        # Add question heading
        q_heading = doc.add_heading(f"Differentiate between {item['topic']}", level=2)
        q_heading.style.font.size = Pt(13)
        doc.add_paragraph()

        # Add Table
        headers = item['headers']
        points = item['points']
        
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Shading Accent 1' # Changed table style for visual difference
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Fill headers
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            # Make header bold
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(12)

        # Fill data
        for point in points:
            row_cells = table.add_row().cells
            for i, val in enumerate(point):
                row_cells[i].text = str(val)

        doc.add_paragraph() # Add space after table

    doc.save(output_filename)
    print(f"Successfully saved {output_filename}")

if __name__ == "__main__":
    # Import Version 2 data from our data parts
    all_data = []
    
    try:
        from data_v2_part1 import data as d1
        all_data.extend(d1)
    except ImportError:
        print("Failed to load v2 part 1")
    
    try:
        from data_v2_part2 import data as d2
        all_data.extend(d2)
    except ImportError:
        print("Failed to load v2 part 2")

    try:
        from data_v2_part3 import data as d3
        all_data.extend(d3)
    except ImportError:
        print("Failed to load v2 part 3")
        
    try:
        from data_v2_part4 import data as d4
        all_data.extend(d4)
    except ImportError:
        print("Failed to load v2 part 4")
        
    try:
        from data_v2_part5 import data as d5
        all_data.extend(d5)
    except ImportError:
        print("Failed to load v2 part 5")
        
    try:
        from data_v2_part6 import data as d6
        all_data.extend(d6)
    except ImportError:
        print("Failed to load v2 part 6")
        
    create_document(all_data)
